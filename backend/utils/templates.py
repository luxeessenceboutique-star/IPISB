"""
Document templates — admin uploads an example file (docx/pdf/image), an LLM
pass runs ONCE at upload time to detect which spans are student-specific
("fields"), and the result is stored on the template row. Every later
generation reuses that stored field map — no LLM call after the first
upload ("detect once, reuse forever").

A field can be detected in one of two shapes:
  - "replace": a real value is already written in the example (e.g. "Jean
    Dupont" after "Nom :") — generation swaps it out for the real student's
    value.
  - "append": only a bare label is present with nothing filled in (e.g.
    "Nom :" with a blank line/space after it, as in a fill-in-the-blank
    form) — generation keeps the label and writes the value right after it.

Detection is deliberately bounded to a fixed vocabulary of fields we can
actually fill from real data (SUPPORTED_FIELDS). Anything else in the
document is left untouched as fixed boilerplate. This keeps the LLM from
inventing fields we have no data for, which would otherwise force us to
either leave the original example person's private info in place (wrong)
or blank out things we didn't need to touch (noisy).
"""
import base64
import io
import json
import logging
import os
import re
import unicodedata
from functools import lru_cache

from docx import Document
from openai import OpenAI
from PIL import Image, ImageDraw, ImageFont
from pypdf import PdfReader, PdfWriter
from reportlab.lib.utils import ImageReader
from reportlab.pdfbase.pdfmetrics import stringWidth
from reportlab.pdfgen import canvas

log = logging.getLogger(__name__)

BASE_URL = os.environ.get("OPENAI_BASE_URL", "https://toknroutertybot.tybotflow.com/")
API_KEY = os.environ.get("OPENAI_API_KEY", "")
MODEL = "gpt-4.1-mini"

CONTENT_TYPE_TO_KIND = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "application/pdf": "pdf",
    "image/jpeg": "image",
    "image/png": "image",
}

SUPPORTED_FIELDS = {
    "student_name": "le nom complet de la personne concernée (nom ET prénom ensemble)",
    "last_name": "le nom de famille SEUL, quand le document sépare « Nom » et « Prénom(s) »",
    "first_name": "le ou les prénoms SEULS, quand ils sont séparés du nom de famille",
    "student_email": "son adresse email",
    "class_name": "sa classe, filière ou promotion",
    "date_of_birth": "sa date de naissance",
    "place_of_birth": "son lieu de naissance",
    "city": "la ville où le document est établi/signé (ex. la ville après « Fait à »)",
    "issue_date": "la date d'émission / de délivrance du document",
    "academic_year": "l'année scolaire concernée (ex. 2025-2026)",
    "establishment_name": "le nom de l'établissement scolaire (école, institut, université)",
    "establishment_address": (
        "l'adresse postale ou la ligne de contact de l'établissement (rue, ville, "
        "téléphone). ATTENTION : le nom d'une faculté ou d'un département n'est "
        "PAS une adresse — c'est other_personal"
    ),
    "verification_code": "un code de référence ou de vérification déjà présent sur le document",
    "cin": "son numéro de carte d'identité nationale (CIN)",
    "matricule": (
        "son numéro d'inscription / matricule / numéro d'étudiant attribué par "
        "l'établissement (ex. numéro Apogée, CNE, code étudiant)"
    ),
    "other_personal": (
        "toute AUTRE donnée spécifique à l'exemple qui ne correspond à aucun champ "
        "ci-dessus : téléphone, niveau d'études, "
        "faculté/département, noms des signataires, et pour un stage TOUTES les "
        "informations du stage (nom de l'entreprise, numéro d'immatriculation/RC, "
        "adresse de l'entreprise, dates et durée du stage, sujet du stage, "
        "encadrants/superviseurs)… Chaque valeur sera effacée du document généré "
        "(remplacée par un blanc à compléter)"
    ),
}

FIELD_LABELS = {
    "student_name": "Nom complet",
    "last_name": "Nom de famille",
    "first_name": "Prénom",
    "student_email": "Email",
    "class_name": "Classe / filière",
    "date_of_birth": "Date de naissance",
    "place_of_birth": "Lieu de naissance",
    "city": "Ville",
    "establishment_phone": "Téléphone de l'établissement",
    "issue_date": "Date d'émission",
    "academic_year": "Année scolaire",
    "establishment_name": "Nom de l'établissement",
    "establishment_address": "Adresse de l'établissement",
    "verification_code": "Code de vérification",
    "cin": "N° CIN",
    "matricule": "Matricule",
    "other_personal": "Donnée de l'exemple (effacée)",
    "photo_zone": "Photo du stagiaire",
    "qr_zone": "QR de vérification",
}


def detect_file_kind(content_type: str) -> str:
    kind = CONTENT_TYPE_TO_KIND.get(content_type)
    if not kind:
        raise ValueError(
            "Type de fichier non supporté. Utilisez un fichier Word (.docx), PDF ou image (JPG/PNG)."
        )
    return kind


def _client() -> OpenAI:
    return OpenAI(api_key=API_KEY, base_url=BASE_URL)


def _strip_markdown_json(raw: str) -> str:
    """Strip markdown code fences and trailing garbage from an LLM JSON
    response. Unlike chatbot.agent._strip_markdown_json (which only handles
    a single {...} object), this tracks whichever bracket the response
    actually opens with, since field detection returns a [...] array."""
    raw = raw.strip()
    if "```json" in raw:
        raw = raw.split("```json", 1)[1]
        if "```" in raw:
            raw = raw.split("```", 1)[0]
    elif "```" in raw:
        raw = raw.split("```", 1)[1]
        if "```" in raw:
            raw = raw.split("```", 1)[0]
    raw = raw.strip()
    if not raw:
        return raw
    open_ch, close_ch = ("[", "]") if raw[0] == "[" else ("{", "}")
    depth = 0
    end_pos = -1
    for i, ch in enumerate(raw):
        if ch == open_ch:
            depth += 1
        elif ch == close_ch:
            depth -= 1
            if depth == 0:
                end_pos = i
                break
    if end_pos != -1:
        raw = raw[: end_pos + 1]
    return raw


def _field_vocab_prompt() -> str:
    return "\n".join(f"- {key}: {desc}" for key, desc in SUPPORTED_FIELDS.items())


# ─────────────────────────────────────────────────────────────────────────
# Text extraction (docx / pdf)
# ─────────────────────────────────────────────────────────────────────────

def _extract_docx_text(file_bytes: bytes) -> str:
    doc = Document(io.BytesIO(file_bytes))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text.strip():
                        parts.append(p.text)
    return "\n".join(parts)


def _extract_pdf_text(file_bytes: bytes) -> str:
    import pdfplumber
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        return "\n".join(page.extract_text() or "" for page in pdf.pages)


# ─────────────────────────────────────────────────────────────────────────
# Jinja-style placeholder templates — {{ variable }} markers mean the
# template is machine-made, so detection can be fully deterministic: map
# each variable name to a field, erase anything we can't map. No LLM at
# all, and text outside the placeholders is guaranteed untouched.
# ─────────────────────────────────────────────────────────────────────────

_JINJA_RE = re.compile(r"\{\{\s*([A-Za-z0-9_]+)\s*\}\}")

_JINJA_EXACT = {
    "nom_etudiant": "last_name",
    "prenom_etudiant": "first_name",
    "nom_prenom": "student_name",
    "nom_complet": "student_name",
    "date_naissance": "date_of_birth",
    "lieu_naissance": "place_of_birth",
    "cin_etudiant": "cin",
    "carte_nationale": "cin",
    "code_massar": "matricule",
    "cne": "matricule",
    "numero_etudiant": "matricule",
    "annee_universitaire": "academic_year",
    "annee_scolaire": "academic_year",
    "filiere_etudes": "class_name",
    "email_etudiant": "student_email",
    "date_edition": "issue_date",
    "date_emission": "issue_date",
    "ville_edition": "city",
    "fait_a": "city",
    "nom_etablissement": "establishment_name",
    "adresse_etablissement": "establishment_address",
    "telephone_etablissement": "establishment_phone",
    "code_verification": "verification_code",
}


def _jinja_field_key(var: str) -> str:
    v = var.lower()
    if v in _JINJA_EXACT:
        return _JINJA_EXACT[v]
    # Fuzzy fallbacks for template authors' naming variations. Order matters:
    # the most specific signal wins (e.g. "lieu_de_naissance" is naissance
    # AND lieu — place beats date).
    if "massar" in v or "cne" in v or "apogee" in v or "matricule" in v:
        return "matricule"
    if "cin" in v:
        return "cin"
    if "naissance" in v:
        return "place_of_birth" if "lieu" in v else "date_of_birth"
    if "annee" in v:
        return "academic_year"
    if "filiere" in v or "classe" in v:
        return "class_name"
    if "email" in v or "mail" in v:
        return "student_email"
    if "ville" in v:
        return "city"
    if "adresse" in v:
        return "establishment_address" if "etablissement" in v else "other_personal"
    if "telephone" in v or "tel" in v:
        return "establishment_phone" if "etablissement" in v else "other_personal"
    if "date" in v:
        return "issue_date"
    if "prenom" in v:
        return "first_name"
    if "nom" in v:
        return "establishment_name" if "etablissement" in v else "last_name"
    # Unknown variable (autorisation number, signataire…) — data we don't
    # hold, erased to a visible blank rather than leaving "{{ ... }}".
    return "other_personal"


def _detect_fields_jinja(text: str) -> list[dict] | None:
    """[{'key', 'mode', 'anchor_text'}, …] for every {{ placeholder }} in the
    text, or None if the document has none (→ LLM detection instead)."""
    matches = list(_JINJA_RE.finditer(text))
    if not matches:
        return None
    fields, seen = [], set()
    for m in matches:
        anchor = m.group(0)
        if anchor in seen:
            continue  # replace-all handles repeats of the same placeholder
        seen.add(anchor)
        # "jinja": the anchor is placeholder markup (often monospace/colored),
        # so generation must NOT imitate its typography — see _sample_font.
        fields.append({"key": _jinja_field_key(m.group(1)), "mode": "replace", "anchor_text": anchor, "jinja": True})
    log.info("Jinja template detected: %d placeholders mapped deterministically", len(fields))
    return fields


# ─────────────────────────────────────────────────────────────────────────
# Field detection — one LLM call per NEW template only
# ─────────────────────────────────────────────────────────────────────────

def _detect_fields_from_text(text: str) -> list[dict]:
    prompt = (
        "Tu analyses un exemple de document administratif — il peut être déjà "
        "rempli pour une personne réelle, OU être un formulaire vierge avec "
        "juste des labels/étiquettes à remplir.\n\n"
        "Repère, dans ce texte, tout passage lié à un des champs suivants "
        "(et rien d'autre) :\n"
        f"{_field_vocab_prompt()}\n\n"
        "Pour chaque champ repéré, deux cas possibles :\n"
        '1. Une vraie valeur est déjà écrite (ex. après "Nom :" il y a déjà un '
        'nom) → {"key": "...", "mode": "replace", "anchor_text": "<la valeur '
        'exacte trouvée, SANS le label>"}\n'
        '2. Seul un label/étiquette est présent, sans valeur écrite après '
        '(ex. juste "Nom :" avec rien après) → {"key": "...", "mode": '
        '"append", "anchor_text": "<le texte exact du label, tel qu\'il '
        'apparaît>"}\n\n'
        'EXEMPLE (formulaire vierge) : pour le texte "Nom :\\nPrénom :\\nNuméro '
        'CIN du tuteur :", la réponse attendue est [{"key": "student_name", '
        '"mode": "append", "anchor_text": "Nom :"}] — "Prénom" (déjà couvert '
        'par student_name) et "Numéro CIN du tuteur" ne correspondent à aucun '
        "champ de la liste, on les ignore dans un formulaire vierge.\n\n"
        "RÈGLES IMPORTANTES pour un document déjà rempli :\n"
        "- Si le nom est séparé en deux (ex. \"Nom : BELHADJ\" puis \"Prénoms : "
        "Karim\"), renvoie la valeur du nom de famille comme last_name et celle "
        "des prénoms comme first_name. N'utilise student_name que pour un nom "
        "complet écrit d'un seul tenant.\n"
        "- La ville où le document est établi (ex. \"Rabat\" dans \"Fait à "
        "Rabat, le …\") → city en mode replace, la ville SEULE sans la date.\n"
        "- Les noms des SIGNATAIRES en bas du document (doyen, directeur, "
        "chef de scolarité, responsable… ex. \"Pr. EL HADJ Mohamed\") → une "
        "entrée other_personal PAR NOM, TOUJOURS — ne les oublie jamais, ce "
        "sont des personnes de l'exemple qui ne doivent pas apparaître dans "
        "le document réutilisé.\n"
        "- Le nom et l'adresse de l'établissement de l'EXEMPLE (même s'il "
        "s'agit d'une autre école ou université) doivent être tagués "
        "establishment_name / establishment_address en mode replace — ils "
        "seront remplacés par ceux de notre établissement.\n"
        "- Toute autre valeur propre à l'exemple → une entrée other_personal "
        "PAR VALEUR : numéro d'étudiant, niveau, faculté/département, "
        "signataires, et pour un stage : entreprise, numéro RC, adresse de "
        "l'entreprise, chaque date du stage, durée, sujet, chaque encadrant. "
        "Sois exhaustif — aucune donnée de l'exemple ne doit survivre dans le "
        "document réutilisé.\n"
        "- Le niveau, cycle ou l'année d'études de l'exemple (ex. \"3ème Année "
        "(Licence Professionnelle)\", \"Licence\", \"Master\", \"1ere année\", "
        "\"2ème année\") → une entrée other_personal PAR VALEUR s'il ne "
        "correspond pas exactement à class_name — ne saute JAMAIS ces valeurs, "
        "même écrites seules après un label comme \"Cycle d'étude\" ou "
        "\"Niveau\".\n"
        "- La filière/spécialité/le domaine d'études de l'exemple (ex. "
        "\"Biologie Médicale\") → class_name en mode replace, JAMAIS "
        "other_personal.\n"
        "- Le lieu de naissance de l'exemple (ou son label seul, ex. \"Lieu "
        "de naissance :\") → place_of_birth, JAMAIS other_personal.\n"
        "- Le nom d'établissement figurant dans l'en-tête du document → "
        "establishment_name, même s'il est écrit en majuscules ou collé sans "
        "espaces (le texte peut provenir d'une reconnaissance optique).\n"
        "- Une valeur qui se répète plusieurs fois dans le document : une "
        "seule entrée suffit, toutes les occurrences seront remplacées "
        "automatiquement.\n\n"
        "N'hésite pas à répondre même si un champ n'est représenté que par un "
        "label seul, sans aucune valeur — c'est un cas fréquent et normal "
        "(formulaire vierge), ne le saute pas par excès de prudence.\n\n"
        "Ignore tout le reste : formules fixes, texte générique. Si un champ "
        "n'apparaît nulle part (ni valeur ni label), ne l'inclus pas.\n\n"
        "Réponds UNIQUEMENT avec un JSON valide : un tableau d'objets comme "
        "décrit ci-dessus.\n\n"
        f"TEXTE DU DOCUMENT :\n{text}\n\nJSON :"
    )
    def one_pass(seed: int) -> list[dict]:
        resp = _client().chat.completions.create(
            model=MODEL,
            messages=[{"role": "user", "content": prompt}],
            temperature=0,
            max_tokens=1500,
            seed=seed,
        )
        original = resp.choices[0].message.content or "[]"
        raw = _strip_markdown_json(original)
        try:
            return json.loads(raw)
        except json.JSONDecodeError:
            log.warning("Field detection JSON parse failed. original=%r stripped=%r", original, raw)
            return []

    # Three passes, merged: the model's recall fluctuates between runs even at
    # temperature 0, and detection happens only once per template — a missed
    # field here stays missed forever. The union recovers what any pass saw.
    data = one_pass(7) + one_pass(1337) + one_pass(42)

    by_anchor: dict[str, dict] = {}
    for f in data:
        if not (f.get("key") in SUPPORTED_FIELDS and f.get("mode") in ("replace", "append") and f.get("anchor_text")):
            continue
        key = f["key"]
        # Known model failure mode: a faculty/department line tagged as the
        # establishment's address. Writing our postal address mid-sentence
        # where a faculty was named produces nonsense — blank it instead.
        if key == "establishment_address":
            t = _norm_text(f["anchor_text"])
            if "facult" in t or "departement" in t or "département" in t:
                key = "other_personal"
        sig = _norm_text(f["anchor_text"]).strip()
        existing = by_anchor.get(sig)
        if existing is None:
            by_anchor[sig] = {"key": key, "mode": f["mode"], "anchor_text": f["anchor_text"]}
        elif existing["key"] == "other_personal" and key != "other_personal":
            # The passes disagree: one recognized a real field where the other
            # just saw "example data". The specific key wins — it fills real
            # student data instead of blanking.
            existing["key"] = key
            existing["mode"] = f["mode"]
    return list(by_anchor.values())


# Unicode dash variants all normalize to a plain hyphen so that a value like
# "Licence Professionnelle - Génie Informatique" (table) still matches its
# body-text repetition "Licence Professionnelle – Génie Informatique".
_DASHES = str.maketrans({c: "-" for c in "‐‑‒–—―"})
_STRIP_PUNCT = ".,;:()«»\"'"


def _norm_text(s: str) -> str:
    """Length-preserving normalization for in-string searching (docx path)."""
    return s.translate(_DASHES).lower()


_ELISION = re.compile(r"^(?:qu|[ldjcnmst])['’]")


def _norm_token(s: str) -> str:
    """Aggressive per-word normalization for PDF word matching. French
    elisions glue onto the next word ("l'Université") and would otherwise
    never match an anchor that starts at "Université"."""
    t = s.translate(_DASHES).lower().strip(_STRIP_PUNCT)
    return _ELISION.sub("", t)


def _find_text_bboxes(page, target: str) -> list[dict]:
    """Locate ALL occurrences of `target` on a pdfplumber page via consecutive
    word matching, tolerant to case, dash variants, and trailing punctuation.
    Matches can span line breaks (words stay consecutive in reading order)."""
    target_words = [_norm_token(w) for w in target.split()]
    target_words = [w for w in target_words if w]
    if not target_words:
        return []
    words = page.extract_words()
    norm_words = [_norm_token(w["text"]) for w in words]
    n = len(target_words)
    hits = []
    i = 0
    while i <= len(words) - n:
        if norm_words[i:i + n] == target_words:
            window = words[i:i + n]
            hits.append({
                "x0": min(w["x0"] for w in window),
                "x1": max(w["x1"] for w in window),
                "top": min(w["top"] for w in window),
                "bottom": max(w["bottom"] for w in window),
                # Per-word rects: a match that wraps across lines has a union
                # bbox spanning both lines, and whiting THAT out would erase
                # neighbouring text. Generation erases word-by-word instead.
                "words": [
                    {"x0": w["x0"], "x1": w["x1"], "top": w["top"], "bottom": w["bottom"]}
                    for w in window
                ],
            })
            i += n
        else:
            i += 1
    return hits


def _is_garbled(text: str) -> bool:
    """True if extraction hit any undecodable glyphs — some PDFs embed fonts
    with a broken/missing ToUnicode map, so accented characters decode to
    U+FFFD. A single occurrence means the font is unreliable (French text is
    full of accents), so we don't wait for a high ratio — clean boilerplate
    elsewhere in the document would dilute a percentage-based check. No text
    library fixes this; falling back to reading the page as a picture
    (vision model) sidesteps it entirely."""
    if len(text.strip()) < 20:
        return True
    return "�" in text


def _vision_detect_fields(png_bytes: bytes, width: int, height: int) -> list[dict]:
    """Shared vision-based detection: works on any rendered page/photo. Reads
    pixels directly, so it's immune to broken PDF font encodings and needs no
    text layer at all (scanned documents included)."""
    b64 = base64.b64encode(png_bytes).decode()
    prompt = (
        "Tu analyses une image d'un exemple de document administratif — il "
        "peut être déjà rempli pour une personne réelle, OU être un "
        f"formulaire vierge avec juste des labels à remplir. L'image fait "
        f"{width}x{height} pixels (origine en haut à gauche).\n\n"
        "Repère les zones liées à un des champs suivants (et rien d'autre) :\n"
        f"{_field_vocab_prompt()}\n\n"
        "Pour chaque champ repéré, deux cas possibles :\n"
        '1. Une vraie valeur est déjà écrite (ex. un nom après "Nom :") → '
        '"mode": "replace", et le bbox doit entourer UNIQUEMENT cette valeur '
        "(pas le label).\n"
        '2. Seul un label est visible, sans valeur écrite → "mode": "append", '
        "et le bbox doit entourer UNIQUEMENT le texte du label lui-même "
        '(ex. juste "Nom :"), PAS l\'espace vide ni la ligne/le trait '
        "entier — la position d'écriture sera calculée automatiquement juste "
        "après ce bbox.\n\n"
        "N'hésite pas à répondre même si un champ n'est représenté que par un "
        "label seul, sans aucune valeur — c'est un cas fréquent et normal "
        "(formulaire vierge), ne le saute pas par excès de prudence.\n\n"
        "Pour un document déjà rempli : le nom/l'adresse de l'établissement de "
        "l'exemple (même une autre école) → establishment_name / "
        "establishment_address en mode replace ; toute autre valeur propre à "
        "la personne de l'exemple (numéro d'étudiant, signataires, entreprise "
        "de stage…) → une entrée other_personal par valeur, mode replace, "
        "elle sera effacée. Si une valeur apparaît plusieurs fois, renvoie "
        "une entrée par occurrence visible.\n\n"
        "Ignore tout le reste. Si un champ n'apparaît nulle part, ne l'inclus "
        "pas.\n\n"
        "Réponds UNIQUEMENT avec un JSON valide : un tableau d'objets "
        '{"key": "<champ>", "mode": "replace"|"append", '
        '"bbox": [x0, y0, x1, y1]} avec des coordonnées pixel exactes et '
        "précises (une boîte serrée autour du seul texte concerné, pas plus "
        "large).\n\nJSON :"
    )
    resp = _client().chat.completions.create(
        model=MODEL,
        messages=[{
            "role": "user",
            "content": [
                {"type": "text", "text": prompt},
                {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{b64}"}},
            ],
        }],
        temperature=0,
        max_tokens=800,
    )
    raw = _strip_markdown_json(resp.choices[0].message.content or "[]")
    data = json.loads(raw)
    fields = []
    for f in data:
        if not (
            f.get("key") in SUPPORTED_FIELDS
            and f.get("mode") in ("replace", "append")
            and isinstance(f.get("bbox"), list)
            and len(f["bbox"]) == 4
        ):
            continue
        x0, y0, x1, y1 = f["bbox"]
        if f["mode"] == "append":
            # bbox is the LABEL's own box — write just after it, same line,
            # rather than trusting the model to gauge "where's the blank
            # space" (a much fuzzier spatial estimate than "where's this text").
            gap = max(6, (x1 - x0) * 0.08)
            write_bbox = [x1 + gap, y0, x1 + gap + 260, y1]
        else:
            write_bbox = [x0, y0, x1, y1]
        fields.append({"key": f["key"], "mode": f["mode"], "bbox": write_bbox})
    return fields


def _suppress_nested_pdf_hits(located: list[dict]) -> list[dict]:
    """Drop replace-hits contained inside a bigger replace-hit on the same
    page. Example: 'RABAT' matched by the city field inside the header line
    'UNIVERSITÉ MOHAMMED V DE RABAT' already covered by establishment_name —
    the big span's replacement wins, the inner one would double-print."""
    def area(b):
        return max(0.0, b["x1"] - b["x0"]) * max(0.0, b["bottom"] - b["top"])

    keep = []
    for i, f in enumerate(located):
        if f["mode"] == "replace":
            bi, ai, nested = f["bbox"], area(f["bbox"]), False
            for j, g in enumerate(located):
                if i == j or g["mode"] != "replace" or g["page"] != f["page"]:
                    continue
                bj = g["bbox"]
                if area(bj) <= ai:
                    continue
                ox = max(0.0, min(bi["x1"], bj["x1"]) - max(bi["x0"], bj["x0"]))
                oy = max(0.0, min(bi["bottom"], bj["bottom"]) - max(bi["top"], bj["top"]))
                if ai > 0 and ox * oy / ai > 0.6:
                    nested = True
                    break
            if nested:
                continue
        keep.append(f)
    return keep


def _detect_fields_pdf_via_text(file_bytes: bytes, text: str) -> list[dict]:
    import pdfplumber
    fields = _detect_fields_jinja(text) or _detect_fields_from_text(text)
    located: list[dict] = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for field in fields:
            hits = []
            for page_num, page in enumerate(pdf.pages):
                for bbox in _find_text_bboxes(page, field["anchor_text"]):
                    hits.append((page_num, bbox))
            if not hits:
                log.info("Anchor not located on any page, dropping: %r", field["anchor_text"])
                continue
            if field["mode"] == "append":
                # A label needs filling only once — at its first occurrence.
                page_num, anchor_bbox = hits[0]
                gap = 4
                located.append({
                    **field,
                    "page": page_num,
                    "bbox": {
                        "x0": anchor_bbox["x1"] + gap,
                        "x1": anchor_bbox["x1"] + gap + 200,
                        "top": anchor_bbox["top"],
                        "bottom": anchor_bbox["bottom"],
                    },
                })
            else:
                # replace: the example person's value must vanish EVERYWHERE
                # it appears (tables, body text, signature lines) — one stored
                # entry per occurrence.
                for page_num, bbox in hits:
                    located.append({**field, "page": page_num, "bbox": bbox})
    return _suppress_nested_pdf_hits(located)


def _detect_fields_pdf_via_vision(file_bytes: bytes) -> list[dict]:
    """Reads each page as a picture instead of extracted text. Needed for
    broken font encodings and scanned pages with no text layer at all.
    OCR provides exact geometry; the vision model's own coordinate guesses
    proved unreliable and are kept only as a last resort."""
    import pdfplumber
    fields: list[dict] = []
    RESOLUTION = 150
    scale = 72 / RESOLUTION  # pdfplumber page geometry is in points at 72 dpi
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page_num, page in enumerate(pdf.pages):
            page_image = page.to_image(resolution=RESOLUTION).original.convert("RGB")
            lines = _ocr_lines(page_image)
            page_fields: list[dict] = []
            if lines:
                text = "\n".join(ln["text"] for ln in lines)
                tagged = _detect_fields_from_text(text)
                page_fields = _locate_fields_via_ocr(lines, tagged)
            if not page_fields:
                buf = io.BytesIO()
                page_image.save(buf, format="PNG")
                page_fields = _vision_detect_fields(buf.getvalue(), page_image.width, page_image.height)
            for f in page_fields:
                x0, y0, x1, y1 = f["bbox"]
                fields.append({
                    **{k: v for k, v in f.items() if k != "bbox"},
                    "page": page_num,
                    "bbox": {"x0": x0 * scale, "x1": x1 * scale, "top": y0 * scale, "bottom": y1 * scale},
                })
    return fields


def _detect_fields_pdf(file_bytes: bytes) -> list[dict]:
    text = _extract_pdf_text(file_bytes)

    if not _is_garbled(text):
        fields = _detect_fields_pdf_via_text(file_bytes, text)
        if fields:
            return fields
        log.info("Text-based PDF detection found nothing — retrying via vision")

    return _detect_fields_pdf_via_vision(file_bytes)


# ─────────────────────────────────────────────────────────────────────────
# OCR-anchored detection for images and scanned PDFs. The vision model's
# pixel-coordinate guesses proved wildly imprecise (wrong scale, wrong rows),
# so OCR supplies exact word geometry and the LLM only decides WHICH text is
# a field — the same division of labour as the born-digital PDF path.
# ─────────────────────────────────────────────────────────────────────────

@lru_cache(maxsize=1)
def _ocr_engine():
    from rapidocr_onnxruntime import RapidOCR
    return RapidOCR()


def _ocr_lines(img: Image.Image) -> list[dict]:
    """Run OCR on a PIL image → [{'text', 'x0', 'x1', 'top', 'bottom'}, …]."""
    import numpy as np
    result, _ = _ocr_engine()(np.array(img))
    lines = []
    for box, text, _score in result or []:
        xs = [p[0] for p in box]
        ys = [p[1] for p in box]
        lines.append({"text": text, "x0": min(xs), "x1": max(xs), "top": min(ys), "bottom": max(ys)})
    return lines


def _norm_loose(s: str) -> str:
    """Accent- and spacing-insensitive form for OCR matching: OCR frequently
    drops spaces ("BiologieMedicale") and accents ("Etablissement")."""
    s = unicodedata.normalize("NFKD", s)
    s = "".join(c for c in s if not unicodedata.combining(c))
    return "".join(c for c in s.lower() if c.isalnum())


def _locate_in_ocr(lines: list[dict], anchor: str) -> list[dict]:
    """All occurrences of `anchor` among OCR lines. Sub-line matches get a
    char-proportional slice of the line box — approximate but close, and
    exact whenever the anchor is a full line (the common case)."""
    na = _norm_loose(anchor)
    if not na:
        return []
    hits = []
    for ln in lines:
        nl = _norm_loose(ln["text"])
        if not nl:
            continue
        idx = nl.find(na)
        while idx != -1:
            w = ln["x1"] - ln["x0"]
            hits.append({
                "x0": ln["x0"] + w * idx / len(nl),
                "x1": ln["x0"] + w * (idx + len(na)) / len(nl),
                "top": ln["top"],
                "bottom": ln["bottom"],
            })
            idx = nl.find(na, idx + len(na))
    return hits


def _suppress_nested_hits(entries: list[dict]) -> list[dict]:
    """Drop a hit that sits inside another field's hit. Happens when one
    anchor is a substring of another line's text (e.g. "Biologie Médicale"
    occurring inside "Institut Supérieur de la Biologie Médicale") — patching
    both would stamp a blank on top of the first replacement."""
    def area(b):
        return max(0.0, b[2] - b[0]) * max(0.0, b[3] - b[1])

    kept: list[dict] = []
    for e in sorted(entries, key=lambda e: -area(e["bbox"])):
        x0, y0, x1, y1 = e["bbox"]
        nested = False
        for k in kept:
            kx0, ky0, kx1, ky1 = k["bbox"]
            ix = max(0.0, min(x1, kx1) - max(x0, kx0))
            iy = max(0.0, min(y1, ky1) - max(y0, ky0))
            if area(e["bbox"]) > 0 and (ix * iy) / area(e["bbox"]) > 0.6:
                nested = True
                break
        if not nested:
            kept.append(e)
    return kept


def _locate_fields_via_ocr(lines: list[dict], fields: list[dict]) -> list[dict]:
    """Anchor LLM-tagged fields onto OCR geometry (image pixel space)."""
    located = []
    for f in fields:
        hits = _locate_in_ocr(lines, f["anchor_text"])
        if not hits:
            log.info("OCR anchor not located, dropping: %r", f["anchor_text"])
            continue
        if f["mode"] == "append":
            h = hits[0]
            gap = max(6, (h["bottom"] - h["top"]) * 0.4)
            located.append({**f, "bbox": [h["x1"] + gap, h["top"], h["x1"] + gap + 260, h["bottom"]]})
        else:
            for h in hits:
                located.append({**f, "bbox": [h["x0"], h["top"], h["x1"], h["bottom"]]})
    return _suppress_nested_hits(located)


def _detect_qr_zone(img: Image.Image) -> list[float] | None:
    """Locate a QR code deterministically via OpenCV's detector."""
    import cv2
    import numpy as np
    bgr = np.array(img)[:, :, ::-1].copy()
    # The Aruco-based detector handles stylized/decorative QR graphics that
    # the classic detector misses.
    for det in (getattr(cv2, "QRCodeDetectorAruco", None), cv2.QRCodeDetector):
        if det is None:
            continue
        try:
            ok, points = det().detect(bgr)
        except Exception:
            continue
        if ok and points is not None:
            pts = points.reshape(-1, 2)
            return [float(pts[:, 0].min()), float(pts[:, 1].min()), float(pts[:, 0].max()), float(pts[:, 1].max())]
    return None


def _detect_photo_zone(img: Image.Image, exclude: list[list[float]]) -> list[float] | None:
    """Locate the ID photo without any face model: a photo is a compact block
    of high LOCAL VARIANCE (busy pixels) whose content spans thousands of
    continuous colours. Flat design elements — logos, icons, headings — are
    busy too, but use only a handful of colours, so a colour-richness test
    separates them cleanly."""
    import cv2
    import numpy as np
    arr = np.array(img)
    gray = np.array(img.convert("L"))

    bs = 8  # block size for the local-variance map
    h2, w2 = (gray.shape[0] // bs) * bs, (gray.shape[1] // bs) * bs
    blocks = gray[:h2, :w2].astype(np.float32)
    blocks = blocks.reshape(h2 // bs, bs, w2 // bs, bs).transpose(0, 2, 1, 3).reshape(h2 // bs, w2 // bs, -1)
    busy = (blocks.std(axis=2) > 10).astype(np.uint8)
    # Mask out known non-photo busy zones (the QR) so a narrow gap can't let
    # the photo region merge with them.
    for ex in exclude:
        x0, y0, x1, y1 = (int(v // bs) for v in ex)
        busy[max(0, y0 - 1):y1 + 2, max(0, x0 - 1):x1 + 2] = 0

    # Colour-band boundaries (e.g. a card's header/body transition) show up as
    # 1-block-thin busy stripes spanning the whole card, welding the photo to
    # unrelated regions. Opening removes thin lines; solid 2-D blobs survive.
    opened = cv2.morphologyEx(busy, cv2.MORPH_OPEN, np.ones((2, 2), np.uint8))

    def color_richness(px0: int, py0: int, px1: int, py1: int) -> int:
        region = arr[py0:py1, px0:px1]
        if region.size == 0:
            return 0
        q = (region.reshape(-1, region.shape[2])[:, :3] // 8).astype(np.uint32)
        return len(np.unique(q[:, 0] * 1024 + q[:, 1] * 32 + q[:, 2]))

    num, labels, stats, _ = cv2.connectedComponentsWithStats(opened, connectivity=8)
    img_area = img.width * img.height
    best: tuple[float, list[int]] | None = None
    for lbl in range(1, num):
        x, y, w, h = (stats[lbl, cv2.CC_STAT_LEFT], stats[lbl, cv2.CC_STAT_TOP],
                      stats[lbl, cv2.CC_STAT_WIDTH], stats[lbl, cv2.CC_STAT_HEIGHT])
        px0, py0, px1, py1 = x * bs, y * bs, (x + w) * bs, (y + h) * bs
        area = (px1 - px0) * (py1 - py0)
        if not (0.005 * img_area <= area <= 0.35 * img_area):
            continue
        aspect = (px1 - px0) / max(1, (py1 - py0))
        if not (0.3 <= aspect <= 2.0):
            continue
        if color_richness(px0, py0, px1, py1) < 300:  # logos/text sit far below; photos in the thousands
            continue
        if best is None or area > best[0]:
            best = (area, [x, y, x + w, y + h])  # block coords
    if best is None:
        return None

    # Refine against the pre-opening map: the opening may have eaten the
    # photo's smooth edges. Within a small window around the seed, keep the
    # rows/columns that are substantially busy — that's the photo rectangle,
    # while a stray 1-block stripe passing through never qualifies.
    sx0, sy0, sx1, sy1 = best[1]
    m = 4
    wx0, wy0 = max(0, sx0 - m), max(0, sy0 - m)
    wx1, wy1 = min(busy.shape[1], sx1 + m), min(busy.shape[0], sy1 + m)
    window = busy[wy0:wy1, wx0:wx1]
    col_need = max(3, int(0.3 * (sy1 - sy0)))
    row_need = max(3, int(0.3 * (sx1 - sx0)))
    cols = np.where(window.sum(axis=0) >= col_need)[0]
    rows = np.where(window.sum(axis=1) >= row_need)[0]
    if len(cols) and len(rows):
        fx0, fx1 = (wx0 + cols.min()) * bs, (wx0 + cols.max() + 1) * bs
        fy0, fy1 = (wy0 + rows.min()) * bs, (wy0 + rows.max() + 1) * bs
    else:
        fx0, fy0, fx1, fy1 = sx0 * bs, sy0 * bs, sx1 * bs, sy1 * bs
    return [float(fx0), float(fy0), float(fx1), float(fy1)]


def _detect_image_zones(img: Image.Image) -> list[dict]:
    """Non-text zones handled at generation time: the ID photo and the QR."""
    zones = []
    qr = _detect_qr_zone(img)
    if qr:
        zones.append({"key": "qr_zone", "mode": "image", "bbox": qr})
    photo = _detect_photo_zone(img, exclude=[qr] if qr else [])
    if photo:
        zones.append({"key": "photo_zone", "mode": "image", "bbox": photo})
    return zones


def _detect_fields_image(file_bytes: bytes, content_type: str) -> list[dict]:
    img = Image.open(io.BytesIO(file_bytes)).convert("RGB")
    zones = _detect_image_zones(img)
    lines = _ocr_lines(img)
    if lines:
        text = "\n".join(ln["text"] for ln in lines)
        fields = _detect_fields_from_text(text)
        located = _locate_fields_via_ocr(lines, fields)
        if located:
            return located + zones
    # No readable text at all — last resort: the vision model's own boxes.
    log.info("Image OCR found nothing usable — falling back to vision bboxes")
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return _vision_detect_fields(buf.getvalue(), img.width, img.height) + zones


def detect_fields(file_bytes: bytes, file_kind: str, content_type: str = "") -> list[dict]:
    """One-time detection pass. Never raises for 'no fields found' — returns []."""
    try:
        if file_kind == "docx":
            text = _extract_docx_text(file_bytes)
            return _detect_fields_jinja(text) or _detect_fields_from_text(text)
        if file_kind == "pdf":
            return _detect_fields_pdf(file_bytes)
        if file_kind == "image":
            return _detect_fields_image(file_bytes, content_type)
    except Exception as e:
        log.warning("Field detection failed (%s): %s", file_kind, e)
        return []
    raise ValueError(f"Unknown file kind: {file_kind}")


# ─────────────────────────────────────────────────────────────────────────
# Generation — pure substitution against the stored field map, no LLM call
# ─────────────────────────────────────────────────────────────────────────

def _resolve_value(key: str, context: dict) -> str:
    """Real data if we have it, otherwise a visible blank rather than ever
    leaving the original example person's info in a different student's doc."""
    value = context.get(key)
    return value if value else "____________"


def _iter_all_paragraphs(doc):
    """Body, tables, and section headers/footers — anchors can live anywhere."""
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p
    for section in doc.sections:
        for part in (section.header, section.footer):
            for p in part.paragraphs:
                yield p
            for table in part.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            yield p


def _replace_in_paragraph(p, anchor: str, replacement: str) -> None:
    """Replace every occurrence of `anchor` (case/dash-insensitive) in a
    paragraph. Word often splits a phrase across formatting runs, so first try
    the cheap per-run replace; if the anchor only exists across run boundaries,
    fall back to rewriting the paragraph text into its first run (loses
    intra-paragraph bold/italic splits for that paragraph — acceptable, since
    the alternative is leaving the example person's data in the document)."""
    n_anchor = _norm_text(anchor)
    if not n_anchor:
        return

    # Pass 1: anchor contained within single runs.
    hit = False
    for run in p.runs:
        n_run = _norm_text(run.text)
        if n_anchor in n_run:
            hit = True
            out, i = "", 0
            while True:
                idx = n_run.find(n_anchor, i)
                if idx == -1:
                    out += run.text[i:]
                    break
                out += run.text[i:idx] + replacement
                i = idx + len(n_anchor)
            run.text = out
    if hit:
        return

    # Pass 2: anchor spans run boundaries.
    full = "".join(r.text for r in p.runs)
    n_full = _norm_text(full)
    if n_anchor not in n_full:
        return
    i = 0
    while True:
        idx = n_full.find(n_anchor, i)
        if idx == -1:
            break
        full = full[:idx] + replacement + full[idx + len(n_anchor):]
        n_full = _norm_text(full)
        i = idx + len(replacement)
    if p.runs:
        p.runs[0].text = full
        for r in p.runs[1:]:
            r.text = ""


def _generate_docx(file_bytes: bytes, fields: list[dict], context: dict) -> bytes:
    doc = Document(io.BytesIO(file_bytes))

    for field in fields:
        anchor = field["anchor_text"]
        new_value = _resolve_value(field["key"], context)
        replacement = new_value if field["mode"] == "replace" else f"{anchor} {new_value}"
        for p in _iter_all_paragraphs(doc):
            _replace_in_paragraph(p, anchor, replacement)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def _map_to_base14(fontname: str) -> str:
    """Closest reportlab built-in to an embedded PDF font name (often a
    subset like 'ABCDEF+Arial-BoldMT'). Exact match is impossible without
    embedding the original font; family + weight + slant gets close enough
    that an inserted value no longer stands out."""
    n = fontname.split("+")[-1].lower()
    bold = any(w in n for w in ("bold", "black", "heavy", "semibold", "demi"))
    italic = "italic" in n or "oblique" in n
    if any(w in n for w in ("times", "serif", "georgia", "garamond", "book")):
        if bold and italic:
            return "Times-BoldItalic"
        return "Times-Bold" if bold else ("Times-Italic" if italic else "Times-Roman")
    if "courier" in n or "mono" in n:
        suffix = "-BoldOblique" if bold and italic else "-Bold" if bold else "-Oblique" if italic else ""
        return "Courier" + suffix
    suffix = "-BoldOblique" if bold and italic else "-Bold" if bold else "-Oblique" if italic else ""
    return "Helvetica" + suffix


def _color_to_rgb(color) -> tuple[float, float, float]:
    if color is None:
        return (0.0, 0.0, 0.0)
    if isinstance(color, (int, float)):
        return (float(color),) * 3
    vals = [float(v) for v in color]
    if len(vals) == 1:
        return (vals[0],) * 3
    if len(vals) == 3:
        return (vals[0], vals[1], vals[2])
    if len(vals) == 4:  # CMYK
        c, m, y, k = vals
        return ((1 - c) * (1 - k), (1 - m) * (1 - k), (1 - y) * (1 - k))
    return (0.0, 0.0, 0.0)


def _sample_font(plumber_page, bbox: dict, exclude_rects: list[dict] | None = None) -> dict | None:
    """Dominant font name / size / fill colour of the template's own
    characters inside `bbox` — or, for append fields whose bbox covers empty
    space, of the characters sharing that text line. `exclude_rects` (jinja
    templates: EVERY placeholder region on the page) keeps `{{ marker }}`
    markup — often monospace/colored, and there may be several on one line —
    out of the sample so only the document's real typography is copied."""
    from collections import Counter

    def chars_in(x0, x1, top, bottom):
        found = []
        for ch in plumber_page.chars:
            cx = (ch["x0"] + ch["x1"]) / 2
            cy = (ch["top"] + ch["bottom"]) / 2
            if exclude_rects and any(
                r["x0"] - 1 <= cx <= r["x1"] + 1 and r["top"] - 1 <= cy <= r["bottom"] + 1
                for r in exclude_rects
            ):
                continue
            if x0 - 1 <= cx <= x1 + 1 and top - 1 <= cy <= bottom + 1:
                found.append(ch)
        return found

    hits = [] if exclude_rects else chars_in(bbox["x0"], bbox["x1"], bbox["top"], bbox["bottom"])
    if not hits:  # blank space or markup anchor — copy the line's font
        hits = chars_in(0, plumber_page.width, bbox["top"], bbox["bottom"])
    if not hits:
        return None
    fontname = Counter(ch.get("fontname") or "" for ch in hits).most_common(1)[0][0]
    sizes = sorted(float(ch.get("size") or 0) for ch in hits)
    size = sizes[len(sizes) // 2]
    def col_key(ch):
        c = ch.get("non_stroking_color")
        return tuple(c) if isinstance(c, (list, tuple)) else (c,)
    color = Counter(col_key(ch) for ch in hits).most_common(1)[0][0]
    rgb = _color_to_rgb(color[0] if len(color) == 1 else color)
    # A near-white fill would make the inserted value invisible on paper —
    # only trust sampled colours that would actually read as ink.
    if min(rgb) > 0.85:
        rgb = (0.0, 0.0, 0.0)
    return {"font": _map_to_base14(fontname), "size": size, "rgb": rgb}


def _generate_pdf(file_bytes: bytes, fields: list[dict], context: dict) -> bytes:
    import pdfplumber

    reader = PdfReader(io.BytesIO(file_bytes))
    writer = PdfWriter()
    plumber = pdfplumber.open(io.BytesIO(file_bytes))

    fields_by_page: dict[int, list[dict]] = {}
    for field in fields:
        fields_by_page.setdefault(field["page"], []).append(field)

    for page_num, page in enumerate(reader.pages):
        page_fields = fields_by_page.get(page_num, [])
        if page_fields:
            width = float(page.mediabox.width)
            height = float(page.mediabox.height)
            # Every rect that gets whited out on this page — characters in
            # these zones no longer exist, so they must not count as
            # obstacles when measuring how much room a new value has.
            erased_rects = [
                w
                for f in page_fields
                if f["mode"] == "replace"
                for w in (f.get("words") or [f["bbox"]])
            ]

            def _avail_width(first_word):
                """Room from the write position to the next character that
                survives generation on the same line. Writing over blank
                space is fine; writing over kept text is what makes a
                document look tampered with."""
                if page_num >= len(plumber.pages):
                    return None
                obstacle = None
                for ch in plumber.pages[page_num].chars:
                    if ch["x0"] <= first_word["x0"] + 1:
                        continue
                    cy = (ch["top"] + ch["bottom"]) / 2
                    if not (first_word["top"] - 1 <= cy <= first_word["bottom"] + 1):
                        continue
                    cx = (ch["x0"] + ch["x1"]) / 2
                    if any(
                        r["x0"] - 1 <= cx <= r["x1"] + 1 and r["top"] - 1 <= cy <= r["bottom"] + 1
                        for r in erased_rects
                    ):
                        continue
                    if obstacle is None or ch["x0"] < obstacle:
                        obstacle = ch["x0"]
                if obstacle is None:
                    return None
                return max(obstacle - first_word["x0"] - 2, 30)

            buf = io.BytesIO()
            c = canvas.Canvas(buf, pagesize=(width, height))
            for field in page_fields:
                bbox = field["bbox"]
                new_value = _resolve_value(field["key"], context)
                # pdfplumber's `top`/`bottom` are measured from the page's TOP;
                # reportlab draws from the BOTTOM — flip the y-axis.
                pad = 1.5
                if field["mode"] == "replace":
                    c.setFillColorRGB(1, 1, 1)
                    # Erase word-by-word: a value wrapped across lines has a
                    # union bbox spanning both lines, and one big white rect
                    # would wipe out neighbouring text.
                    for w in field.get("words") or [bbox]:
                        c.rect(
                            w["x0"] - pad,
                            height - w["bottom"] - pad,
                            (w["x1"] - w["x0"]) + 2 * pad,
                            (w["bottom"] - w["top"]) + 2 * pad,
                            fill=1, stroke=0,
                        )
                first = (field.get("words") or [bbox])[0]
                # Imitate the template's own typography at this exact spot so
                # the inserted value doesn't read as an insertion.
                sampled = None
                if page_num < len(plumber.pages):
                    try:
                        sampled = _sample_font(plumber.pages[page_num], bbox, exclude_rects=erased_rects if field.get("jinja") else None)
                    except Exception:
                        sampled = None
                if sampled:
                    font_name = sampled["font"]
                    font_size = max(6.5, min(30, sampled["size"]))
                    rgb = sampled["rgb"]
                else:
                    line_h = first["bottom"] - first["top"]
                    font_name = "Helvetica"
                    font_size = max(8, min(12, line_h * 0.85))
                    rgb = (0, 0, 0)
                # Shrink to fit the actual free space so a long value neither
                # spills over kept text (measured obstacle) nor over table
                # borders / off the page (fallback slot width).
                avail = _avail_width(first)
                if avail is None:
                    avail = max(bbox["x1"] - bbox["x0"], 160)
                while font_size > 6.5 and stringWidth(new_value, font_name, font_size) > avail:
                    font_size -= 0.5
                c.setFillColorRGB(*rgb)
                c.setFont(font_name, font_size)
                c.drawString(first["x0"], height - first["bottom"] + pad, new_value)
            c.save()
            buf.seek(0)
            overlay = PdfReader(buf).pages[0]
            page.merge_page(overlay)
        writer.add_page(page)

    plumber.close()
    out = io.BytesIO()
    writer.write(out)
    return out.getvalue()


def _find_system_font(size: int) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    for path in (
        "C:\\Windows\\Fonts\\arial.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/System/Library/Fonts/Supplemental/Arial.ttf",
    ):
        try:
            return ImageFont.truetype(path, size)
        except Exception:
            continue
    return ImageFont.load_default()


def _sample_bg_color(img: Image.Image, x0: int, y0: int, x1: int, y1: int) -> tuple[int, int, int]:
    """Median colour of a thin ring just outside the rect — what the patch
    should look like once the old text is gone. White would glow on a card's
    coloured background."""
    import numpy as np
    arr = np.array(img)
    h, w = arr.shape[:2]
    pads = []
    top_a, top_b = max(0, y0 - 6), max(0, y0 - 2)
    if top_b > top_a:
        pads.append(arr[top_a:top_b, max(0, x0):min(w, x1)])
    bot_a, bot_b = min(h, y1 + 2), min(h, y1 + 6)
    if bot_b > bot_a:
        pads.append(arr[bot_a:bot_b, max(0, x0):min(w, x1)])
    if not pads:
        return (255, 255, 255)
    ring = np.concatenate([p.reshape(-1, arr.shape[2]) for p in pads if p.size])
    if ring.size == 0:
        return (255, 255, 255)
    med = np.median(ring, axis=0).astype(int)
    return (int(med[0]), int(med[1]), int(med[2]))


def _fetch_photo(url: str) -> Image.Image | None:
    import httpx
    try:
        resp = httpx.get(url, timeout=15, follow_redirects=True)
        resp.raise_for_status()
        return Image.open(io.BytesIO(resp.content)).convert("RGB")
    except Exception as e:
        log.warning("Student photo fetch failed (%s): %s", url, e)
        return None


def _paste_cover(base: Image.Image, photo: Image.Image, x0: int, y0: int, x1: int, y1: int) -> None:
    """Scale-and-center-crop the photo to exactly fill the zone (no bars, no
    distortion) — like CSS object-fit: cover."""
    w, h = max(1, x1 - x0), max(1, y1 - y0)
    scale = max(w / photo.width, h / photo.height)
    rw, rh = int(photo.width * scale + 0.5), int(photo.height * scale + 0.5)
    resized = photo.resize((rw, rh), Image.LANCZOS)
    cx, cy = (rw - w) // 2, (rh - h) // 2
    base.paste(resized.crop((cx, cy, cx + w, cy + h)), (x0, y0))


def _apply_image_zones(img: Image.Image, zones: list[dict], context: dict) -> None:
    import qrcode
    for zone in zones:
        x0, y0, x1, y1 = (int(v) for v in zone["bbox"])
        if zone["key"] == "photo_zone":
            photo = _fetch_photo(context["student_photo_url"]) if context.get("student_photo_url") else None
            if photo is not None:
                _paste_cover(img, photo, x0, y0, x1, y1)
            else:
                # No photo on file: a neutral placeholder beats keeping the
                # example person's face on another student's card.
                ImageDraw.Draw(img).rectangle([x0, y0, x1, y1], fill=(229, 231, 235))
        elif zone["key"] == "qr_zone" and context.get("verify_url"):
            # The detector returns the QR symbol bounds; the printed graphic
            # extends a few px further (quiet zone). White out a padded area
            # first so no modules of the old QR peek around the new one.
            pad = max(4, int((x1 - x0) * 0.08))
            ImageDraw.Draw(img).rectangle(
                [max(0, x0 - pad), max(0, y0 - pad), min(img.width, x1 + pad), min(img.height, y1 + pad)],
                fill=(255, 255, 255),
            )
            qr_img = qrcode.make(context["verify_url"]).convert("RGB")
            side = min(x1 - x0, y1 - y0)
            cx = x0 + ((x1 - x0) - side) // 2
            cy = y0 + ((y1 - y0) - side) // 2
            _paste_cover(img, qr_img, cx, cy, cx + side, cy + side)


def _generate_image_as_pdf(file_bytes: bytes, fields: list[dict], context: dict) -> bytes:
    img = Image.open(io.BytesIO(file_bytes)).convert("RGB")

    _apply_image_zones(img, [f for f in fields if f.get("mode") == "image"], context)

    draw = ImageDraw.Draw(img)
    for field in fields:
        if field.get("mode") == "image":
            continue
        x0, y0, x1, y1 = (int(v) for v in field["bbox"])
        new_value = _resolve_value(field["key"], context)
        bg = _sample_bg_color(img, x0, y0, x1, y1)
        if field["mode"] == "replace":
            draw.rectangle([x0 - 2, y0 - 2, x1 + 2, y1 + 2], fill=bg)
        # Dark ink on light backgrounds, white on dark ones (card headers).
        luminance = 0.299 * bg[0] + 0.587 * bg[1] + 0.114 * bg[2]
        color = (255, 255, 255) if luminance < 128 else (25, 35, 80)
        # Fit: start from the erased line's height, shrink until the value
        # fits the space to the right edge of the card.
        avail = max(x1 - x0, min(260, img.width - x0 - 8))
        font_size = max(10, int((y1 - y0) * 0.85))
        font = _find_system_font(font_size)
        while font_size > 9 and draw.textlength(new_value, font=font) > avail:
            font_size -= 1
            font = _find_system_font(font_size)
        # Vertically centre in the original line box.
        draw.text((x0, y0 + max(0, ((y1 - y0) - font_size * 1.2) / 2)), new_value, fill=color, font=font)

    # Wrap the flattened image into a single-page PDF for a consistent
    # preview/download experience with every other generated document.
    img_buf = io.BytesIO()
    img.save(img_buf, format="PNG")
    img_buf.seek(0)

    pdf_buf = io.BytesIO()
    width_pt, height_pt = img.width * 72 / 150, img.height * 72 / 150
    c = canvas.Canvas(pdf_buf, pagesize=(width_pt, height_pt))
    c.drawImage(ImageReader(img_buf), 0, 0, width=width_pt, height=height_pt)
    c.save()
    return pdf_buf.getvalue()


def generate_from_template(file_bytes: bytes, file_kind: str, fields: list[dict], context: dict) -> tuple[bytes, str]:
    """Returns (output_bytes, output_content_type). No LLM call — pure substitution."""
    if file_kind == "docx":
        return (
            _generate_docx(file_bytes, fields, context),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    if file_kind == "pdf":
        return _generate_pdf(file_bytes, fields, context), "application/pdf"
    if file_kind == "image":
        return _generate_image_as_pdf(file_bytes, fields, context), "application/pdf"
    raise ValueError(f"Unknown file kind: {file_kind}")
