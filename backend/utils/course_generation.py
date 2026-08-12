import io
import json
import logging
import os
import re

import pdfplumber
from docx import Document as DocxDocument
from langchain_core.messages import HumanMessage, SystemMessage
from langchain_openai import ChatOpenAI
from supabase import Client

log = logging.getLogger(__name__)

BASE_URL = os.environ.get("OPENAI_BASE_URL", "https://toknroutertybot.tybotflow.com/")
API_KEY  = os.environ.get("OPENAI_API_KEY", "")
MODEL    = "gpt-4.1-mini"

LIBRARY_BUCKET = "reference-library"

# Extensions we can text-extract today. Legacy .doc (pre-2007 binary Word) is
# excluded on purpose — python-docx only reads OOXML (.docx) — so any
# CDC/programme still in .doc format is silently skipped here, not crashed
# on. Convert it to .docx once and it'll be picked up automatically.
_EXTRACTABLE = {".docx", ".pdf"}

# Generic French words filtered out of relevance keywords so they don't match
# every document ever ("cours", "chapitre"...) and drown out the words that
# actually identify a subject.
_STOPWORDS_FR = {
    "cours", "chapitre", "module", "cette", "avec", "pour", "dans",
    "les", "des", "une", "leur", "leurs", "sont", "cela", "ainsi", "notion",
    "notions", "elements", "element", "introduction", "generale", "general",
}


def _client(temperature: float = 0.4) -> ChatOpenAI:
    return ChatOpenAI(model=MODEL, base_url=BASE_URL, api_key=API_KEY, temperature=temperature)


def _extract_text(data: bytes, filename: str) -> str | None:
    ext = os.path.splitext(filename)[1].lower()
    try:
        if ext == ".docx":
            doc = DocxDocument(io.BytesIO(data))
            parts = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))
            return "\n".join(parts)
        if ext == ".pdf":
            with pdfplumber.open(io.BytesIO(data)) as pdf:
                return "\n".join(page.extract_text() or "" for page in pdf.pages)
    except Exception as e:
        log.warning("course_generation: failed to extract %s: %s", filename, e)
        return None
    return None  # unsupported extension (.doc, .xls, .xlsx for now)


def _keywords_from(*texts: str | None) -> set[str]:
    """Words (4+ letters, French accents included) pulled from a course/module's
    own title+description/objectives — used only as a soft relevance signal
    in gather_grounding_text below, never as a filter."""
    words: set[str] = set()
    for t in texts:
        if not t:
            continue
        for w in re.findall(r"[a-zàâäéèêëïîôöùûüç]{4,}", t.lower()):
            if w not in _STOPWORDS_FR:
                words.add(w)
    return words


def _sort_by_relevance(rows: list[dict], keywords: set[str]) -> list[dict]:
    """Reorders (never drops) library rows so ones whose title matches the
    requested subject are read — and so survive the char_budget cutoff —
    before generic ones. No keywords, or no title matches any of them: order
    is left exactly as it came in (upload order), so this is a no-op for
    every caller that doesn't opt in."""
    if not keywords:
        return rows

    def score(row: dict) -> int:
        title = (row.get("title") or "").lower()
        return sum(title.count(k) for k in keywords)

    return sorted(rows, key=score, reverse=True)  # stable: ties keep upload order


def _relevance_excerpt(text: str, keywords: set[str], budget: int) -> str:
    """Within ONE document too long to include whole, keep the paragraphs
    that actually mention the requested subject instead of blindly taking
    the head — a single CDC/programme file often covers many UF at once, so
    'the first `budget` characters' can easily land on the wrong one.
    Falls back to plain head-truncation (today's behaviour) whenever there's
    nothing to key off, or nothing in the text matches — this only ever
    changes WHICH characters are kept, never whether a document is used."""
    if len(text) <= budget or not keywords:
        return text[:budget]
    paragraphs = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]
    if not paragraphs:
        return text[:budget]

    scored = [(sum(p.lower().count(k) for k in keywords), i, p) for i, p in enumerate(paragraphs)]
    if max(s for s, _, _ in scored) == 0:
        return text[:budget]  # nothing in this document matches — old behaviour

    chosen: set[int] = set()
    used = 0
    for s, i, p in sorted(scored, key=lambda x: -x[0]):
        if s == 0 or used >= budget:
            break
        chosen.add(i)
        used += len(p) + 2
    # keep the selected paragraphs in their original reading order
    return "\n\n".join(p for i, p in enumerate(paragraphs) if i in chosen)[:budget]


def gather_grounding_text(
    db: Client,
    specialty_id: str | None,
    categories: list[str],
    char_budget: int = 12000,
    relevance_hint: str | None = None,
) -> str:
    """Pull matching library_files (this filière's docs + any 'général' ones
    in the same categories), download from storage, extract text, and
    concatenate up to char_budget, each chunk clearly labelled with its
    source title.

    `relevance_hint`: free text (course/module title + description/objectives)
    used to prioritize which documents get read first, and which passages of
    an over-long one get kept — see _sort_by_relevance/_relevance_excerpt.
    Purely a re-ranking: with no hint, or no keyword match anywhere, this
    function's output is byte-for-byte what it was before that existed
    (upload order, head-truncation) — nothing that worked before can start
    returning less than it used to."""
    keywords = _keywords_from(relevance_hint)

    query = db.from_("library_files").select("*").in_("category", categories)
    if specialty_id:
        query = query.or_(f"specialty_id.eq.{specialty_id},specialty_id.is.null")
    else:
        query = query.is_("specialty_id", "null")
    rows = query.order("created_at").execute().data or []
    rows = _sort_by_relevance(rows, keywords)

    chunks: list[str] = []
    used = 0
    skipped: list[str] = []
    for row in rows:
        if used >= char_budget:
            break
        ext = os.path.splitext(row["file_path"])[1].lower()
        if ext not in _EXTRACTABLE:
            skipped.append(row["title"])
            continue
        try:
            data = db.storage.from_(LIBRARY_BUCKET).download(row["file_path"])
        except Exception as e:
            log.warning("course_generation: failed to download %s: %s", row["file_path"], e)
            continue
        text = _extract_text(data, row["file_path"])
        if not text:
            continue
        remaining = char_budget - used
        text = _relevance_excerpt(text, keywords, remaining)
        chunk = f"--- Source : {row['title']} ({row['category']}) ---\n{text}"
        chunks.append(chunk)
        used += len(chunk)

    if skipped:
        log.info("course_generation: skipped non-extractable sources: %s", skipped)
    return "\n\n".join(chunks)


def _extract_json(text: str, expect: str = "array") -> object:
    pattern = r"\[.*\]" if expect == "array" else r"\{.*\}"
    match = re.search(pattern, text, re.DOTALL)
    if not match:
        raise ValueError(f"No JSON {expect} found in model output")
    return json.loads(match.group(0))


def generate_module_outline(
    course_title: str, filiere_label: str, grounding: str, already_published: list[str] | None = None,
    description: str | None = None,
) -> list[dict]:
    """Stage 1: break ONE existing course/UF (e.g. 'Droit législation') into
    a sequence of teaching chapters — objectives + TH/TP hour split — grounded
    in that filière's CDC/programme/exam-bank material. Returns drafts only;
    caller persists as draft course_modules rows.

    `already_published`: titles of chapters already live for students. When
    given, the model extends the plan with what's still missing instead of
    re-covering the same ground — without this, regenerating after a partial
    publish produces a second "Introduction" chapter with no idea the first
    one already exists.

    `description`: the professor's own 2-3 sentence brief of what this UF
    actually covers (required on new courses — see routers/courses.py). This
    is what keeps a vague/ambiguous title from being outweighed by whatever
    reference material happened to be retrieved: the model is told in plain
    terms what it's structuring, not left to infer it from the title alone."""
    description_block = f"Description du cours fournie par le formateur : {description}\n" if description else ""
    if already_published:
        published_block = (
            "\nCHAPITRES DÉJÀ PUBLIÉS POUR CE COURS (ne les recrée pas, ne les recouvre pas — génère "
            "UNIQUEMENT les chapitres qui manquent encore pour compléter le programme) :\n"
            + "\n".join(f"- {t}" for t in already_published) + "\n"
        )
    else:
        published_block = ""

    system = (
        "Tu es un ingénieur pédagogique dans un institut de formation paramédicale privé au Maroc, "
        "qui structure le contenu interne d'une unité de formation (UF) déjà définie dans le "
        "programme officiel, en s'appuyant sur les documents de référence fournis (cahier des "
        "charges, programme de formation, banque de questions d'examen).\n\n"
        f"UF à structurer : « {course_title} » — filière {filiere_label}.\n"
        f"{description_block}"
        f"{published_block}\n"
        "DOCUMENTS DE RÉFÉRENCE :\n"
        f"{grounding}\n\n"
        "Découpe cette UF en séquences pédagogiques (chapitres) cohérentes qui couvrent, au minimum, "
        "TOUS les sujets mentionnés dans les questions d'examen ci-dessus — ce sont les compétences "
        "que les stagiaires doivent maîtriser.\n\n"
        "RÈGLES STRICTES :\n"
        "- Le plan doit correspondre AU TITRE ET À LA DESCRIPTION DU COURS ci-dessus en priorité — "
        "si le contenu d'un document de référence ne correspond manifestement pas à ce sujet, "
        "ignore ce document plutôt que de produire un plan hors-sujet.\n"
        "- Réponds UNIQUEMENT avec un tableau JSON valide, sans texte avant/après, sans balises markdown.\n"
        "- Chaque élément : {\"title\": str, \"objectives\": str (1-2 phrases), "
        "\"hours_theory\": number, \"hours_practice\": number}.\n"
        "- 3 à 8 chapitres, dans un ordre pédagogique logique (notions de base avant notions avancées).\n"
        "- hours_practice à 0 si la matière est purement théorique (ex. droit).\n"
        "- Français clair, registre professionnel adapté à des formateurs.\n"
    )
    resp = _client(temperature=0.3).invoke([SystemMessage(content=system), HumanMessage(content="Génère le plan.")])
    raw = resp.content if isinstance(resp.content, str) else str(resp.content)
    items = _extract_json(raw, "array")

    cleaned = []
    for it in items:
        if not it.get("title"):
            continue
        cleaned.append({
            "title": str(it["title"]),
            "objectives": str(it.get("objectives") or ""),
            "hours_theory": float(it.get("hours_theory") or 0),
            "hours_practice": float(it.get("hours_practice") or 0),
        })
    if not cleaned:
        raise ValueError("La génération IA n'a produit aucun chapitre valide. Réessayez.")
    return cleaned


def generate_lesson_content(course_title: str, module_title: str, objectives: str, grounding: str) -> str:
    """Stage 2: full narrative lesson text (markdown) for ONE chapter,
    grounded in the same filière's reference material. This is a DRAFT —
    the caller must keep it unpublished until a qualified formateur has
    reviewed factual claims, especially anything legal/clinical.

    Deliberately pushed toward depth over brevity: this feeds a PDF read by
    students as their actual course material, not a slide deck — a chapter
    that's mostly bullet points reads as a summary of a course, not the
    course itself."""
    system = (
        "Tu es un formateur expérimenté et un expert du domaine, qui rédige un chapitre de cours "
        "complet et approfondi pour des stagiaires en formation paramédicale au Maroc (niveau "
        "avancé — l'équivalent d'un support de cours universitaire, pas une fiche de révision).\n\n"
        f"UF : « {course_title} »\n"
        f"Chapitre à rédiger : « {module_title} »\n"
        f"Objectif pédagogique visé : {objectives}\n\n"
        "DOCUMENTS DE RÉFÉRENCE (source prioritaire quand ils couvrent le sujet ; complète avec tes "
        "connaissances expertes du domaine pour tout ce qu'ils ne couvrent pas — le chapitre doit "
        "être complet même là où la documentation source est courte) :\n"
        f"{grounding}\n\n"
        "EXIGENCES DE FOND — c'est le critère le plus important :\n"
        "- L'unité de base du texte est le PARAGRAPHE rédigé : phrases complètes, articulées, qui "
        "expliquent le POURQUOI et le COMMENT (mécanismes, causes, conséquences, nuances, "
        "exceptions, liens entre notions) — pas seulement le QUOI.\n"
        "- Les listes à puces restent utiles et bienvenues, mais seulement pour ce qui EST "
        "réellement une liste : une classification, les étapes d'une procédure, des critères "
        "à cocher, une énumération d'exemples. N'utilise jamais une liste pour éviter de rédiger un "
        "paragraphe explicatif — le mélange naturel prose/liste d'un vrai manuel est l'objectif, pas "
        "l'élimination totale des listes.\n"
        "- Ajoute, aux endroits pertinents, les remarques qu'un formateur expérimenté ferait à "
        "l'oral et qu'on ne trouve pas dans un simple résumé : pièges et confusions fréquentes "
        "(« Attention à ne pas confondre... »), points qui reviennent souvent à l'examen, "
        "renvois à d'autres notions/chapitres du programme, mise en contexte marocain concrète "
        "(texte de loi, institution, pratique professionnelle réelle). Intègre-les dans le fil du "
        "texte (pas une liste séparée de 'astuces').\n"
        "- Illustre systématiquement avec des exemples concrets et des cas d'application "
        "professionnels ancrés dans le contexte marocain et dans la pratique du métier visé par "
        "cette filière.\n"
        "- Couvre le sujet en profondeur et intégralement au regard de l'objectif visé : ne laisse "
        "aucun sous-thème pertinent de côté sous prétexte qu'il n'apparaît pas dans les documents "
        "de référence.\n"
        "- Registre soutenu, vocabulaire technique précis et correctement défini à sa première "
        "occurrence — ce texte doit pouvoir servir de référence, pas seulement d'introduction.\n"
        "- Longueur : un chapitre de ce niveau fait normalement plusieurs centaines de mots par "
        "sous-section ; un chapitre entier en dessous de 900 mots est presque certainement trop "
        "superficiel pour cet objectif.\n\n"
        "SCHÉMAS : quand une section du chapitre EST fondamentalement une classification/typologie "
        "(ex. « les 3 types de X »), en plus de l'expliquer en prose, représente-la AUSSI comme un "
        "schéma visuel, avec ce format exact (bloc de code de langage `diagram`, 2 à 4 catégories, "
        "3 à 6 éléments courts par catégorie) :\n"
        "```diagram\n"
        "TITLE: <titre du schéma>\n"
        "CATEGORIE: <nom catégorie 1>\n"
        "- <élément court>\n"
        "- <élément court>\n"
        "CATEGORIE: <nom catégorie 2>\n"
        "- <élément court>\n"
        "```\n"
        "N'utilise ce format que pour une vraie classification/typologie centrale au chapitre (pas "
        "plus d'un ou deux schémas par chapitre) — jamais pour du texte qui devrait être un "
        "paragraphe.\n\n"
        "FORME : Markdown structuré (## et ### pour les sections, **gras** pour les termes clés à "
        "leur définition, tableaux pour les comparaisons/synthèses, blocs ```diagram comme décrit "
        "ci-dessus le cas échéant). Termine par une synthèse qui relie les notions entre elles (pas "
        "une simple liste récapitulative). "
        "Ne réponds qu'avec le contenu du cours, sans texte d'introduction ni commentaire méta.\n"
    )
    resp = _client(temperature=0.5).invoke([SystemMessage(content=system), HumanMessage(content="Rédige le chapitre, en profondeur.")])
    content = resp.content if isinstance(resp.content, str) else str(resp.content)
    content = content.strip()
    if not content:
        raise ValueError("La génération IA n'a produit aucun contenu. Réessayez.")
    return content


def edit_text_with_ai(text: str, instruction: str) -> str:
    """Phase 7 (spec §17) — rewrites ONE slide text element's content given
    a teacher's free-text instruction ('simplifie ça', 'ajoute un exemple
    marocain', ...). Deliberately plain-text output: unlike
    generate_lesson_content's markdown chapter prose, a slide text element
    renders its content literally (frontend/.../editor.$moduleId.tsx's
    Konva.Text and the read-only SlideDeckViewer both just print `content`
    as-is), so markdown syntax here would show up as literal asterisks and
    hashes on the slide."""
    system = (
        "Tu es un assistant pédagogique qui aide un formateur à modifier le texte d'UN élément "
        "d'une diapositive de cours (formation paramédicale au Maroc), selon son instruction.\n\n"
        "RÈGLES STRICTES :\n"
        "- Réponds UNIQUEMENT avec le nouveau texte, sans aucun commentaire, explication ni texte "
        "d'introduction, et sans guillemets englobants.\n"
        "- Pas de syntaxe Markdown (pas de **gras**, pas de #titres, pas de blocs de code) — c'est "
        "du texte brut affiché tel quel sur une diapositive. Des puces simples ('- ') sont "
        "acceptées si le contenu s'y prête.\n"
        "- Reste concis : c'est un élément de diapositive, pas un paragraphe de cours complet — "
        "quelques phrases ou quelques puces courtes, adapté à ce qui tient visuellement sur une "
        "diapositive.\n"
        "- Registre professionnel, français clair, adapté à des stagiaires en formation "
        "paramédicale au Maroc.\n"
    )
    human = f"Texte actuel de la diapositive :\n\"\"\"\n{text}\n\"\"\"\n\nInstruction : {instruction}"
    resp = _client(temperature=0.5).invoke([SystemMessage(content=system), HumanMessage(content=human)])
    result = resp.content if isinstance(resp.content, str) else str(resp.content)
    result = result.strip().strip('"').strip()
    if not result:
        raise ValueError("La génération IA n'a produit aucun texte. Réessayez.")
    return result


_SLIDE_ROLES = {"title", "subtitle", "body", "bullets"}


def regenerate_slide_with_ai(texts: list[str], instruction: str) -> list[dict]:
    """Whole-slide AI editing (spec §17, extended beyond one element) — the
    teacher selects a SLIDE, not a single text box, and the AI restructures
    ALL the text on it at once ('transforme en exercice', 'transforme en
    liste à puces', ...), which single-element editing can't do since it
    can only rewrite what's already there, not add/remove/reshape blocks.

    Deliberately returns ROLES + content, never x/y/width/height — asking
    an LLM for pixel coordinates invites overlapping or off-canvas
    elements. The caller lays each role out with the same fixed positions
    the slide templates already use, so results are always visually sane;
    the teacher can still drag/resize afterward like anything else.
    Images/shapes already on the slide are left untouched by the caller —
    this function only ever sees/returns text."""
    numbered = "\n---\n".join(texts) if texts else "(diapositive actuellement sans texte)"
    system = (
        "Tu es un assistant pédagogique qui aide un formateur à restructurer le contenu textuel "
        "d'UNE diapositive de cours (formation paramédicale au Maroc), selon son instruction.\n\n"
        f"Contenu textuel actuel de la diapositive, dans l'ordre :\n{numbered}\n\n"
        f"Instruction : {instruction}\n\n"
        "RÈGLES STRICTES :\n"
        "- Réponds UNIQUEMENT avec un tableau JSON valide, sans texte avant/après, sans balises markdown.\n"
        "- Chaque élément : {\"role\": \"title\"|\"subtitle\"|\"body\"|\"bullets\", \"content\": str}.\n"
        "- \"title\" : au maximum UN par diapositive, court (moins de 8 mots), pas de ponctuation finale.\n"
        "- \"subtitle\" : optionnel, une phrase courte.\n"
        "- \"body\" : un paragraphe de texte brut, sans Markdown.\n"
        "- \"bullets\" : plusieurs lignes commençant par '- ', séparées par des retours à la ligne, "
        "dans le MÊME champ content (une liste à puces = un seul élément \"bullets\").\n"
        "- Entre 1 et 5 éléments au total, adaptés à ce qui tient visuellement sur une diapositive.\n"
        "- Registre professionnel, français clair, adapté à des stagiaires en formation paramédicale.\n"
    )
    resp = _client(temperature=0.5).invoke([SystemMessage(content=system), HumanMessage(content="Restructure la diapositive.")])
    raw = resp.content if isinstance(resp.content, str) else str(resp.content)
    items = _extract_json(raw, "array")

    cleaned = []
    for it in items:
        content = str(it.get("content") or "").strip()
        if not content:
            continue
        role = it.get("role") if it.get("role") in _SLIDE_ROLES else "body"
        cleaned.append({"role": role, "content": content})
    if not cleaned:
        raise ValueError("La génération IA n'a produit aucun contenu. Réessayez.")
    return cleaned
